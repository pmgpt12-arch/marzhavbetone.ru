import os
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
header_font = Font(color="FFFFFF", bold=True, size=11)


def create_excel(filepath, sheet_name="Лист1", headers=None, rows=None, col_widths=None):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet_name
    
    if headers:
        ws.append(headers)
        for col_idx, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = thin_border
    
    if rows:
        for r in rows:
            ws.append(r)
            row_num = ws.max_row
            for col_idx in range(1, len(r) + 1):
                cell = ws.cell(row=row_num, column=col_idx)
                cell.border = thin_border
                cell.alignment = Alignment(vertical="top", wrap_text=True)
    
    if col_widths:
        for idx, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(idx)].width = w
    
    wb.save(filepath)
    print(f"  [Excel] {filepath}")


def create_word_doc(filepath, title, sections):
    doc = Document()
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    for section in sections:
        if "heading" in section:
            doc.add_heading(section["heading"], level=1)
        
        if "text" in section:
            doc.add_paragraph(section["text"])
        
        if "bullet" in section:
            for item in section["bullet"]:
                doc.add_paragraph(item, style='List Bullet')
        
        if "table" in section:
            table_data = section["table"]
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row_data in enumerate(table_data):
                    for j, cell_text in enumerate(row_data):
                        table.cell(i, j).text = str(cell_text)
                doc.add_paragraph()
        
        if "numbered" in section:
            for item in section["numbered"]:
                doc.add_paragraph(item, style='List Number')
    
    doc.save(filepath)
    print(f"  [Word]  {filepath}")


def create_pdf_simple(filepath, title, lines):
    c = canvas.Canvas(filepath, pagesize=A4)
    width, height = A4
    
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width/2, height - 2*cm, title)
    
    c.setFont("Helvetica", 11)
    y = height - 3.5*cm
    for line in lines:
        if y < 2*cm:
            c.showPage()
            y = height - 2*cm
            c.setFont("Helvetica", 11)
        c.drawString(2*cm, y, line)
        y -= 0.6*cm
    
    c.save()
    print(f"  [PDF]   {filepath}")


# ==================== FREE 02 ====================
def build_free_02():
    folder = os.path.join(BASE_DIR, "00-free-dopy-ne-v-podarok")
    print(f"\n[BUILD] {folder}")
    
    create_word_doc(
        os.path.join(folder, "01-checklist-do-nachala-doprabot.docx"),
        "ЧЕК-ЛИСТ: Что проверить ДО начала допработ",
        [
            {"heading": "Перед началом любых допработ ответьте на 7 вопросов",
             "numbered": [
                 "Есть ли письменное поручение заказчика на допобъем? (приказ, письмо, протокол совещания)",
                 "Подписано ли соглашение об изменении объемов / сроков / стоимости?",
                 "Есть ли письменное подтверждение цены за единицу допработ?",
                 "Фиксированы ли сроки выполнения допработ?",
                 "Понятно, кто оплачивает материалы для допработ?",
                 "Есть ли запись в журнале работ о допобъеме?",
                 "Сделаны ли фотофиксации 'до' начала допработ?"
             ]},
            {"heading": "Если ответ 'НЕТ' хотя бы на один вопрос",
             "text": "Риск неоплаты допработ резко возрастает. Рекомендуем остановиться и зафиксировать условия в письменном виде. Используйте схему фиксации поручения и шаблон уведомления из этого комплекта."},
            {"heading": "⚠️ Что этот материал НЕ закрывает",
             "bullet": [
                 "Составление юридически корректного допсоглашения с расчетом стоимости",
                 "Пересмотр календарного плана с учетом допработ и простоев",
                 "Работа с претензиями заказчика о качестве допработ",
                 "Взыскание неоплаченных допработ через суд (особенно при отсутствии письменных документов)",
                 "Фотофиксация: полный регламент с привязкой к актам скрытых работ"
             ]},
            {"heading": "🔗 Переход к полному комплекту",
             "text": "Полный комплект 'Допы не в подарок' включает: готовые допсоглашения (объем, сроки, цена), калькулятор стоимости, алгоритм согласования с заказчиком, шаблоны писем о приостановке работ.\n\nСтоимость: 14 900 - 19 900 ₽\nСсылка: [здесь будет ссылка на платный продукт]"},
        ]
    )
    
    create_word_doc(
        os.path.join(folder, "02-shema-fiksacii-porucheniya.docx"),
        "СХЕМА: Как зафиксировать устное поручение на допработу",
        [
            {"heading": "Ситуация",
             "text": "Заказчик устно просит доделать лестничную клетку, увеличить площадь отмостки или заменить материал. Вы начинаете работу - а потом заказчик говорит: 'Я такого не говорил' или 'Это в договоре включено'."},
            {"heading": "Алгоритм фиксации (выполняйте СРАЗУ)",
             "numbered": [
                 "ПЕРЕПИСКА. Напишите в мессенджер / email: 'Подтверждаю получение поручения на [что именно]. Ожидаю письменное подтверждение сроков и цены. Без подтверждения работы не начинаю.' Сохраните скриншоты.",
                 "СЛУЖЕБНАЯ ЗАПИСКА. В тот же день составьте служебную записку руководителю объекта с описанием: кто, когда, что просил, в каком присутствии. Подпишите у руководителя.",
                 "ЖУРНАЛ РАБОТ. Внесите запись в общий журнал работ: 'Получено устное поручение от [ФИО] на [описание работ]. Ожидается письменное подтверждение.'",
                 "ФОТОФИКСАЦИЯ. Сделайте фото состояния объекта ДО начала допработ с привязкой к дате (включите геолокацию на телефоне).",
                 "УВЕДОМЛЕНИЕ. Направьте заказчику официальное письмо с уведомлением о получении поручения (шаблон - в этом комплекте)."
             ]},
            {"heading": "Что НЕЛЬЗЯ делать",
             "bullet": [
                 "Начинать работу без письменного подтверждения стоимости",
                 "Принимать устные обещания 'потом все подпишем'",
                 "Работать 'в надежде' - надежда не является доказательством в суде",
                 "Отправлять КС-2 по допработам до подписания соглашения об изменении цены"
             ]},
            {"heading": "Когда устное поручение уже выполнено, а документов нет",
             "text": "Ситуация сложная, но не безнадежная. В полном комплекте есть алгоритм восстановления документальной базы: через переписку, свидетельские показания, экспертизу объемов."},
            {"heading": "🔗 Переход к полному комплекту",
             "text": "Полный комплект включает: готовые допсоглашения, калькулятор стоимости, регламент фотофиксации, алгоритм согласования, шаблоны писем о приостановке.\n\nСтоимость: 14 900 - 19 900 ₽\nСсылка: [здесь будет ссылка на платный продукт]"},
        ]
    )
    
    create_word_doc(
        os.path.join(folder, "03-shablon-uvedomleniya.docx"),
        "ШАБЛОН УВЕДОМЛЕНИЯ о получении поручения на дополнительные работы",
        [
            {"text": "[На бланке организации]"},
            {"text": ""},
            {"text": "Исх. № _____ от '___' ___________ 20___ г."},
            {"text": ""},
            {"text": "[Должность, ФИО представителя заказчика]"},
            {"text": "[Наименование заказчика]"},
            {"text": ""},
            {"text": "От [Наименование подрядчика]"},
            {"text": ""},
            {"text": "УВЕДОМЛЕНИЕ"},
            {"text": "о получении поручения на выполнение дополнительных работ"},
            {"text": ""},
            {"text": "Уважаемый [ФИО]!"},
            {"text": ""},
            {"text": "Настоящим уведомляем, что [дата, время] в ходе [совещания / телефонного разговора / переписки] нам было сообщено о необходимости выполнения дополнительных работ в рамках договора подряда № [номер] от [дата] (далее - Договор)."},
            {"text": ""},
            {"text": "Содержание поручения:"},
            {"bullet": [
                "[Описание работ]",
                "[Локация / этаж / участок]",
                "[Предполагаемые объемы]"
            ]},
            {"text": ""},
            {"text": "В соответствии с п. [номер] Договора изменение объемов работ, сроков и цены допускается только по соглашению сторон в письменной форме."},
            {"text": ""},
            {"text": "Просим в срок до [дата] направить нам письменное подтверждение поручения с указанием:"},
            {"bullet": [
                "конкретных объемов дополнительных работ,",
                "сроков выполнения,",
                "утвержденной цены за единицу работ,",
                "порядка оплаты."
            ]},
            {"text": ""},
            {"text": "До получения письменного соглашения о дополнительных работах мы не вправе приступить к их выполнению, поскольку это может привести к невозможности согласования стоимости и сроков."},
            {"text": ""},
            {"text": "В случае неполучения подтверждения в указанный срок оставляем за собой право отказаться от выполнения указанных работ без последствий по договору."},
            {"text": ""},
            {"text": "Генеральный директор _________________ [ФИО]"},
            {"text": "М.П."},
            {"text": ""},
            {"text": ""},
            {"text": "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━"},
            {"heading": "💡 Как использовать этот шаблон",
             "numbered": [
                 "Отправьте сразу после получения устного поручения (в тот же день).",
                 "Отправляйте заказным письмом с уведомлением о вручении.",
                 "Сохраняйте квитанцию - это доказательство, что вы действовали добросовестно.",
                 "Если заказчик не отвечает - работы не начинайте. Это ваше право и защита.",
                 "Для срочных работ (авария, безопасность) - фотофиксируйте и уведомляйте параллельно."
             ]},
            {"heading": "🔗 Переход к полному комплекту",
             "text": "Полный комплект включает: готовые допсоглашения, калькулятор стоимости, регламент фотофиксации, алгоритм согласования.\n\nСтоимость: 14 900 - 19 900 ₽\nСсылка: [здесь будет ссылка на платный продукт]"},
        ]
    )
    
    create_excel(
        os.path.join(folder, "04-reestr-doprabot.xlsx"),
        "Реестр допработ",
        headers=["№ п/п", "Дата поручения", "Описание работ", "Объем", "Ед. изм.", "Цена за ед.", "Сумма", "Дата согласования", "Статус", "Примечание"],
        rows=[
            [1, "", "Пример: Устройство отмостки", "150", "м2", "850", "=G2*D2", "", "Согласовано", ""],
            [2, "", "", "", "", "", "", "", "", ""],
            [3, "", "", "", "", "", "", "", "", ""],
            ["", "", "", "", "", "ИТОГО:", "=СУММ(G2:G10)", "", "", ""],
        ],
        col_widths=[8, 14, 30, 10, 10, 12, 12, 16, 14, 20]
    )


if __name__ == "__main__":
    build_free_02()
    print("\nГОТОВО: free-02")
