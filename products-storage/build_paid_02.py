"""Дополнение платного комплекта «Допы не в подарок» (02-dopraboty-bez-poter).

Зачем скрипт существует. Бесплатный материал обещает «за границей
бесплатного» готовые допсоглашения (объём, сроки, цена), и то же обещание
стоит в описании платного манифеста — а в составе комплекта ни одного
допсоглашения не было (D-02). По решению владельца от 14.08.2026 такой
разрыв — дефект продукта и чинится дополнением состава; прецедент —
`11-trebovanie-o-vozvrate-uderzhaniya.docx` в комплекте удержаний.

Обещанную тройку «объём, сроки, цена» закрывают две формы — по образцу
пары из комплекта договора (03-dogovor-podryada), где объём и цена идут
одним соглашением: они и меняются вместе.

Повторный прогон перезаписывает те же два файла тем же содержимым.
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Комплект и его вложенная копия внутри полного комплекта ПТО: копия
# сверяется по содержимому проверкой tools/check_packages.py, поэтому
# файлы кладутся в обе папки одним прогоном.
TARGETS = [
    os.path.join(BASE_DIR, "02-dopraboty-bez-poter"),
    os.path.join(BASE_DIR, "04-polnyy-komplekt-pto",
                 "01-30-bazovye-pakety", "02-dopraboty-bez-poter"),
]

HINT = "Заполните поля в квадратных скобках и удалите подсказки перед отправкой."

# По нормативной сверке 15.08.2026 (ст. 452 и 160 ГК РФ): форма — это
# содержательная часть допсоглашения, и это названо в ней прямо, чтобы
# чек-лист о подписании не спорил с отсутствием подписных блоков.
SCOPE_NOTE = ("Форма — содержательная часть допсоглашения: преамбулу со "
              "сторонами, номер и дату договора подряда и подписные блоки "
              "добавьте по вашему договору.")

CLOSING = ("Остальные условия договора не изменяются. Соглашение вступает "
           "в силу с момента подписания сторонами.")

# Контрольная таблица — та же, что в остальных формах линейки.
CONTROL_TABLE = [
    ["Статус", "Проверка"],
    ["☐", "Сверено с договором и приложениями"],
    ["☐", "Проверены полномочия подписанта и адресат"],
    ["☐", "Приложения названы и приложены"],
    ["☐", "Зафиксирован доказуемый канал отправки"],
    ["☐", "Юридически значимые формулировки проверены специалистом"],
]

DISCLAIMER = ("Важно: шаблон не заменяет анализ договора, проектной "
              "документации и обстоятельств конкретного объекта.")


def create_word_doc(filepath, title, sections):
    doc = Document()
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for section in sections:
        if "heading" in section:
            doc.add_heading(section["heading"], level=1)
        if "text" in section:
            doc.add_paragraph(section["text"])
        if "table" in section:
            table_data = section["table"]
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = 'Table Grid'
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    table.cell(i, j).text = str(cell_text)
            doc.add_paragraph()

    doc.save(filepath)
    print(f"  [Word]  {filepath}")


DOCS = [
    ("11-dopsoglashenie-obem-i-cena.docx",
     "Дополнительное соглашение об изменении объёма и цены "
     "(дополнительные работы)",
     [
         {"text": HINT},
         {"text": SCOPE_NOTE},
         {"heading": "Основание",
          "text": "Поручение на дополнительные работы: [___]\n"
                  "Чем зафиксировано (уведомление, запись в журнале, "
                  "переписка): [___]"},
         {"heading": "Изменение",
          "text": "Добавляемые работы (наименование, единица, "
                  "количество): [___]\n"
                  "Изменение цены договора на: [___]\n"
                  "Порядок и срок оплаты дополнительных работ: [___]"},
         {"heading": "Документы",
          "text": "Ведомость объёмов дополнительных работ и расчёт их "
                  "стоимости являются приложениями."},
         {"heading": "Заключительные положения",
          "text": CLOSING},
         {"heading": "Контроль перед использованием",
          "table": CONTROL_TABLE},
         {"text": DISCLAIMER},
     ]),
    ("12-dopsoglashenie-sroki-doprabot.docx",
     "Дополнительное соглашение об изменении сроков "
     "(дополнительные работы)",
     [
         {"text": HINT},
         {"text": SCOPE_NOTE},
         {"heading": "Основание",
          "text": "Дополнительные работы по допсоглашению № [___] "
                  "от [___]\n"
                  "Влияние на выполнение работ (затронутые операции): "
                  "[___]"},
         {"heading": "Новые сроки",
          "text": "Затронутые этапы и даты: [___]\n"
                  "Новая дата завершения работ: [___]\n"
                  "Обновлённый календарный план является приложением."},
         {"heading": "Заключительные положения",
          "text": CLOSING},
         {"heading": "Контроль перед использованием",
          "table": CONTROL_TABLE},
         {"text": DISCLAIMER},
     ]),
]


def build_paid_02():
    for target in TARGETS:
        if not os.path.isdir(target):
            raise SystemExit(f"нет папки {target} — состав не тот, "
                             "что ожидался")
        for filename, title, sections in DOCS:
            create_word_doc(os.path.join(target, filename), title, sections)


if __name__ == "__main__":
    build_paid_02()
